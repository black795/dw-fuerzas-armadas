import time
import pyautogui
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report_with_screenshot():
    print("Esperando 10 segundos para tomar la captura...")
    for i in range(10, 0, -1):
        print(f"{i}...")
        time.sleep(1)
        
    print("¡Tomando captura de pantalla!")
    screenshot_path = "evidencia_azure.png"
    # Tomar captura real de la pantalla
    pyautogui.screenshot(screenshot_path)
    
    print("Generando el documento Word...")
    doc = Document()
    
    # Titulo
    title = doc.add_heading('Informe de Evaluación 3: Data Warehouse Fuerzas Armadas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datos del estudiante
    doc.add_paragraph('Estudiante: Alan Perez Medrano')
    doc.add_paragraph('Carnet de Identidad: 5756067')
    doc.add_paragraph('Compañero de Trabajo: Esteban Medrano')
    
    # Resumen paso a paso
    doc.add_heading('Resumen del Procedimiento de Solución', level=1)
    p = doc.add_paragraph()
    p.add_run('A continuación se detalla paso a paso el procedimiento técnico llevado a cabo para la resolución del requerimiento institucional de las Fuerzas Armadas de Bolivia:\n').bold = True
    
    doc.add_paragraph('1. Extracción y Respaldo de Datos Crudos: Se diseñó un script en Python que conectó a la base de datos origen (erp-fuerzas-armadas-bolivia-2), extrayendo las 10 tablas transaccionales en formato CSV. Este proceso sirve para salvaguardar la data histórica en la capa local.', style='List Number')
    doc.add_paragraph('2. Aprovisionamiento de Infraestructura en la Nube (IaC): Se implementó código de Terraform configurando todos los servicios en la nube de Microsoft Azure. Este proceso creó de forma automatizada un Grupo de Recursos (gr-sisger-dwffaa-5756067), una cuenta de Storage Gen2 (saucbdwffaa5756067) con contenedores Bronze y Silver, un Servidor SQL (sql-ucb-sisger-ffaa-5756067) y un Azure Data Factory.', style='List Number')
    doc.add_paragraph('3. Modelado Dimensional (Data Warehouse): Empleando los diseños proporcionados, se ejecutó código SQL DDL directamente sobre el nuevo servidor SQL en Azure para desplegar la capa Gold. Se crearon las tablas de hechos y dimensiones que soportarán el sistema analítico.', style='List Number')
    
    doc.add_heading('Evidencia Fotográfica de la Infraestructura en Azure', level=1)
    doc.add_paragraph('La siguiente imagen captura el portal de Microsoft Azure donde se aprecian los recursos desplegados de forma exitosa mediante nuestra automatización de Terraform, validando la autoría mediante la inclusión del carnet de identidad.')
    
    # Insertar la captura
    doc.add_picture(screenshot_path, width=Inches(6.0))
    
    # Guardar
    doc.save('Informe_Alan_Perez.docx')
    print("¡Documento guardado exitosamente!")

if __name__ == '__main__':
    create_report_with_screenshot()
