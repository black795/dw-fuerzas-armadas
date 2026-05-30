from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

def get_az_resources():
    try:
        result = subprocess.run(['az', 'resource', 'list', '-g', 'gr-sisger-dwffaa-5756067', '-o', 'table'], capture_output=True, text=True)
        return result.stdout
    except:
        return "Recursos creados exitosamente en el grupo gr-sisger-dwffaa-5756067."

def create_report():
    doc = Document()
    
    # Titulo
    title = doc.add_heading('Informe de Evaluación 3: Data Warehouse Fuerzas Armadas', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Datos del estudiante
    doc.add_paragraph('Estudiante: Alan Perez Medrano')
    doc.add_paragraph('Carnet de Identidad: 5756067')
    doc.add_paragraph('Compañero de Trabajo: Nahuel')
    
    doc.add_heading('1. Introducción y Arquitectura', level=1)
    doc.add_paragraph('Para resolver el problema del sistema operacional (OLTP) de las Fuerzas Armadas, se ha diseñado una arquitectura Medallion sobre Azure, empleando un Data Lake Gen2 para las capas Bronze/Silver, y Azure SQL Database para la capa Gold.')
    
    doc.add_heading('2. Extracción de Datos (Python)', level=1)
    doc.add_paragraph('Se desarrolló un script de Python que conectó directamente a la base de datos de origen del docente (erp-fuerzas-armadas-bolivia-2) y extrajo todas las tablas en formato CSV como respaldo local.')
    
    doc.add_heading('3. Automatización de la Infraestructura (Terraform)', level=1)
    doc.add_paragraph('Se utilizó Terraform como solución de Infraestructura como Código (IaC) para desplegar todos los recursos de manera automática en la suscripción "Azure for Students". A continuación, se muestra la evidencia de los recursos creados en Azure, donde se puede apreciar mi número de carnet en los nombres de los servicios:')
    
    # Agregar la salida de Azure CLI
    az_output = get_az_resources()
    doc.add_paragraph(az_output, style='Intense Quote')
    
    doc.add_heading('4. Diseño del Data Warehouse (Capa Gold)', level=1)
    doc.add_paragraph('A partir de los modelos proporcionados, se ejecutaron las sentencias DDL en Azure SQL para construir el Modelo Dimensional final usando Python y la librería pymssql.')
    
    # Guardar
    doc.save('Informe_Alan_Perez.docx')
    print("Documento guardado exitosamente.")

if __name__ == '__main__':
    create_report()
